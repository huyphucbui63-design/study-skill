from __future__ import annotations

import json
import re
import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from PIL import Image
from docx import Document
from docx.shared import Mm
from pypdf import PdfReader
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
SKILL = ROOT / "skills" / "kaoyan-memorization-builder"
sys.path.insert(0, str(SKILL / "scripts"))

from project_model import to_print_manifest, validate_project  # noqa: E402


def make_project(source_path: Path, *, status: str = "confirmed", needs_review: bool = False) -> dict:
    return {
        "schema_version": "1.0",
        "title": "极限与连续",
        "subject": "高等数学",
        "scope": "第一章",
        "sources": [{"id": "s1", "source_order": 1, "type": "user_text", "label": "讲义.txt", "path": str(source_path)}],
        "chapters": [{"id": "c1", "source_order": 1, "title": "第一章 极限", "points": [
            {
                "id": "p1", "source_order": 1, "title": "极限定义", "kind": "definition",
                "grading": {
                    "importance": "A",
                    "importance_status": status,
                    "importance_evidence": [{"kind": "user_designation", "statement": "用户指定为核心", "confidence": 1.0}],
                    "personal_weak": True,
                    "weakness_status": status,
                    "weakness_evidence": [{"kind": "user_designation", "statement": "用户指定为薄弱点", "confidence": 1.0}],
                },
                "segments": [{"origin": "source_text", "content": "设函数在某邻域有定义。", "verbatim": True, "needs_review": needs_review, "references": [{"source_id": "s1", "text_range": "1-2"}]},
                             {"origin": "ai_memory_aid", "content": "先看邻域，再看趋近。", "verbatim": False, "needs_review": False, "references": []}],
                "visuals": [{"type": "comparison", "title": "相邻概念", "origin": "ai_summary", "needs_review": False, "references": [{"source_id": "s1", "text_range": "1-2"}], "headers": ["概念", "关注点"], "rows": [["极限", "趋近"], ["连续", "取值"]]}],
                "recall_lines": 2,
            },
            {
                "id": "p2", "source_order": 2, "title": "过程", "kind": "mixed",
                "grading": {
                    "importance": "C",
                    "importance_status": status,
                    "importance_evidence": [{"kind": "source_emphasis", "statement": "原资料小节标题", "confidence": 0.8}],
                    "personal_weak": False,
                    "weakness_status": status,
                    "weakness_evidence": [{"kind": "user_designation", "statement": "用户确认当前不是个人薄弱点", "confidence": 1.0}],
                },
                "segments": [{"origin": "user_note", "content": "补充：注意左右极限。", "verbatim": True, "needs_review": False, "references": []}],
                "visuals": [{"type": "process", "title": "判断顺序", "origin": "user_note", "needs_review": False, "references": [], "items": ["确认定义域", "分别计算左右极限"]}],
            },
        ]}],
        "transformations": [],
        "review": {"status": "confirmed" if status == "confirmed" else "draft", "confirmed_at": "2026-08-07T00:00:00+08:00", "confirmed_by": "user" if status == "confirmed" else "not_confirmed"},
    }


class MemorizationBuilderTests(unittest.TestCase):
    def test_skill_metadata_contract(self) -> None:
        content = (SKILL / "SKILL.md").read_text(encoding="utf-8")
        match = re.match(r"^---\n(.*?)\n---", content, re.DOTALL)
        self.assertIsNotNone(match)
        entries = {}
        for line in match.group(1).splitlines():
            key, value = line.split(":", 1)
            entries[key.strip()] = value.strip()
        self.assertEqual(set(entries), {"name", "description"})
        self.assertEqual(entries["name"], "kaoyan-memorization-builder")
        self.assertRegex(entries["name"], r"^[a-z0-9-]+$")
        self.assertLessEqual(len(entries["description"]), 1024)
        self.assertNotRegex(entries["description"], r"[<>]")
        agent_metadata = (SKILL / "agents" / "openai.yaml").read_text(encoding="utf-8")
        self.assertIn("$kaoyan-memorization-builder", agent_metadata)

    def test_schema_declares_core_contract(self) -> None:
        schema = json.loads((SKILL / "schemas" / "memorization-project.schema.json").read_text(encoding="utf-8"))
        self.assertEqual(schema["$schema"], "https://json-schema.org/draft/2020-12/schema")
        self.assertEqual(schema["properties"]["schema_version"]["const"], "1.0")
        self.assertIn("source_order", schema["$defs"]["source"]["required"])
        self.assertEqual(schema["$defs"]["grading"]["properties"]["importance"]["enum"], ["A", "B", "C"])
        self.assertIn("personal_weak", schema["$defs"]["grading"]["required"])
        self.assertIn("importance_status", schema["$defs"]["grading"]["required"])
        self.assertIn("weakness_status", schema["$defs"]["grading"]["required"])
        self.assertIn("target", schema["$defs"]["transformation"]["required"])
        self.assertEqual(schema["$defs"]["segment"]["properties"]["origin"]["enum"], ["source_text", "ai_summary", "ai_memory_aid", "ai_example", "user_note"])
        self.assertIn("origin", schema["$defs"]["visual"]["required"])

    def test_rejects_unconfirmed_grading_and_fragments(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "source.txt"
            source.write_text("source", encoding="utf-8")
            project = make_project(source, status="ai_suggestion")
            with self.assertRaises(ValueError):
                validate_project(project, require_confirmed=True)
            project = make_project(source, needs_review=True)
            with self.assertRaises(ValueError):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["chapters"][0]["points"][0]["grading"]["weakness_evidence"] = [{"kind": "ai_inference", "statement": "Model guess", "confidence": 0.7}]
            with self.assertRaisesRegex(ValueError, "confirmed R marker"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["chapters"][0]["points"][0]["grading"]["importance_evidence"][0]["confidence"] = 1.5
            with self.assertRaisesRegex(ValueError, "Confidence"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["chapters"][0]["points"][0]["grading"]["weakness_status"] = "ai_suggestion"
            with self.assertRaisesRegex(ValueError, "Unconfirmed R weakness"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["unexpected_top_level"] = True
            with self.assertRaisesRegex(ValueError, "Unexpected field"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["chapters"][0]["points"][0]["id"] = "bad id"
            with self.assertRaisesRegex(ValueError, "Invalid ID"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            del project["chapters"][0]["points"][0]["visuals"][0]["origin"]
            with self.assertRaisesRegex(ValueError, "Visual origin"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["chapters"][0]["points"][0]["visuals"][0]["references"][0]["source_id"] = "missing"
            with self.assertRaisesRegex(ValueError, "Unknown source_id"):
                validate_project(project, require_confirmed=True)

    def test_reorder_requires_explicit_user_authorization(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "source.txt"
            source.write_text("source", encoding="utf-8")
            project = make_project(source)
            project["chapters"][0]["points"].reverse()
            with self.assertRaisesRegex(ValueError, "Source order changed"):
                validate_project(project, require_confirmed=True)
            project["transformations"] = [{"type": "reorder", "target": "chapter:c1", "authorized_by": "user", "authorized_at": "2026-08-07T00:00:00+08:00", "details": "用户要求先复习过程。"}]
            validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["chapters"][0]["points"].reverse()
            project["transformations"] = [{"type": "reorder", "target": "sources", "authorized_by": "user", "authorized_at": "2026-08-07T00:00:00+08:00", "details": "仅授权来源重排。"}]
            with self.assertRaisesRegex(ValueError, "Source order changed"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["transformations"] = [{"type": "reorder", "target": "chapter:missing", "authorized_by": "user", "authorized_at": "2026-08-07T00:00:00+08:00", "details": "错误目标。"}]
            with self.assertRaisesRegex(ValueError, "unknown chapter"):
                validate_project(project, require_confirmed=True)
            project = make_project(source)
            project["transformations"] = [{"type": "reorder", "target": "chapter:c1", "authorized_by": "user", "authorized_at": "not-a-date", "details": "错误时间。"}]
            with self.assertRaisesRegex(ValueError, "Invalid date-time"):
                validate_project(project, require_confirmed=True)

    def test_extracts_synthetic_pdf_and_routes_image_to_review(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            pdf = root / "source.pdf"
            drawing = canvas.Canvas(str(pdf))
            drawing.drawString(72, 760, "Limit definition")
            drawing.save()
            image = root / "formula.png"
            Image.new("RGB", (320, 160), "white").save(image)
            docx = root / "source.docx"
            document = Document()
            document.add_paragraph("Before table")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Term"
            table.cell(0, 1).text = "Meaning"
            document.add_paragraph("After table")
            document.save(docx)
            output = root / "draft.json"
            command = [sys.executable, str(SKILL / "scripts" / "extract_sources.py"), "--title", "Draft", "--subject", "Math", "--output", str(output), str(pdf), str(image), str(docx)]
            result = subprocess.run(command, capture_output=True, text=True)
            self.assertEqual(result.returncode, 0, result.stderr)
            project = json.loads(output.read_text(encoding="utf-8"))
            self.assertEqual([source["type"] for source in project["sources"]], ["pdf", "image", "docx"])
            pdf_segment = project["chapters"][0]["points"][0]["segments"][0]
            self.assertEqual(pdf_segment["references"][0]["pdf_page"], 1)
            image_segment = project["chapters"][0]["points"][1]["segments"][0]
            self.assertTrue(image_segment["needs_review"])
            self.assertIn("待人工识别", image_segment["content"])
            docx_segments = [point["segments"][0] for point in project["chapters"][0]["points"][2:]]
            self.assertEqual([segment["content"] for segment in docx_segments], ["Before table", "Term\tMeaning", "After table"])
            self.assertEqual(docx_segments[1]["references"][0]["text_range"], "表格 1")
            self.assertEqual(project["review"]["status"], "draft")
            grading = project["chapters"][0]["points"][0]["grading"]
            self.assertEqual(grading["importance_status"], "ai_suggestion")
            self.assertEqual(grading["weakness_status"], "ai_suggestion")

    def test_preserves_direct_user_text_verbatim(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            output = Path(folder) / "draft.json"
            original = "精确定义：f(x) 在 x0 连续。"
            result = subprocess.run([sys.executable, str(SKILL / "scripts" / "extract_sources.py"), "--title", "Draft", "--subject", "Math", "--output", str(output), "--text", original], capture_output=True, text=True)
            self.assertEqual(result.returncode, 0, result.stderr)
            project = json.loads(output.read_text(encoding="utf-8"))
            segment = project["chapters"][0]["points"][0]["segments"][0]
            self.assertEqual(segment["content"], original)
            self.assertTrue(segment["verbatim"])
            self.assertFalse(segment["needs_review"])

    def test_conversion_preserves_order_and_independent_markers(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            source = Path(folder) / "source.txt"
            source.write_text("source", encoding="utf-8")
            project = make_project(source)
            manifest = to_print_manifest(project, "bw")
            self.assertEqual([section["title"] for section in manifest["sections"]], ["第一章 极限"])
            points = [block for block in manifest["sections"][0]["blocks"] if block["type"] == "knowledge_point"]
            self.assertEqual([point["title"] for point in points], ["极限定义", "过程"])
            self.assertEqual(points[0]["importance"], "A")
            self.assertTrue(points[0]["personal_weak"])
            self.assertEqual(points[1]["importance"], "C")
            self.assertFalse(points[1]["personal_weak"])
            self.assertIn("重要度依据", points[0]["grading_evidence"])
            self.assertIn("R 依据", points[0]["grading_evidence"])
            comparison = next(block for block in manifest["sections"][0]["blocks"] if block["type"] == "comparison")
            self.assertEqual(comparison["origin_label"], "AI 概括")
            self.assertIn("讲义.txt", comparison["source"])
            project["density"] = "compact"
            self.assertEqual(to_print_manifest(project, "color")["density"], "compact")

    def test_end_to_end_generates_both_profiles(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            source = root / "source.txt"
            source.write_text("source", encoding="utf-8")
            figure = root / "figure.png"
            Image.new("RGB", (320, 160), "white").save(figure)
            project_path = root / "project.json"
            project = make_project(source)
            project["chapters"][0]["points"][0]["visuals"].append({"type": "image", "title": "必要图形", "origin": "source_text", "needs_review": False, "references": [{"source_id": "s1", "text_range": "figure"}], "path": "figure.png", "caption": "定义示意图"})
            second = json.loads(json.dumps(project["chapters"][0], ensure_ascii=False))
            second["id"] = "c2"
            second["source_order"] = 2
            second["title"] = "第二章 连续"
            for index, point in enumerate(second["points"], 1):
                point["id"] = f"p2-{index}"
            project["chapters"].append(second)
            project_path.write_text(json.dumps(project, ensure_ascii=False), encoding="utf-8")
            output = root / "out"
            command = [sys.executable, str(SKILL / "scripts" / "build_memorization.py"), str(project_path), "--output-dir", str(output)]
            result = subprocess.run(command, capture_output=True, text=True)
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertEqual(sorted(path.name for path in output.iterdir()), ["bw-study.docx", "bw-study.pdf", "color-study.docx", "color-study.pdf"])
            color_pdf_text = [page.extract_text() or "" for page in PdfReader(str(output / "color-study.pdf")).pages]
            bw_pdf_text = [page.extract_text() or "" for page in PdfReader(str(output / "bw-study.pdf")).pages]
            self.assertEqual(color_pdf_text, bw_pdf_text)
            qa = subprocess.run([sys.executable, str(SKILL / "scripts" / "qa_material.py"), str(output / "bw-study.pdf"), "--required", "目录", "--required", "第一章 极限", "--required", "第二章 连续", "--required", "[A]", "--required", "[R]", "--toc-entry", "第一章 极限=2", "--toc-entry", "第二章 连续=3", "--min-images", "2", "--min-image-dpi", "150"], capture_output=True, text=True)
            self.assertEqual(qa.returncode, 0, qa.stdout + qa.stderr)
            pdf_report = json.loads(qa.stdout)
            self.assertFalse(pdf_report["toc_mismatches"])
            self.assertGreaterEqual(min(item["effective_dpi"] for item in pdf_report["image_placements"]), 150)
            docx_qa = subprocess.run([sys.executable, str(SKILL / "scripts" / "qa_docx.py"), str(output / "color-study.docx"), str(output / "bw-study.docx"), "--require-toc", "--min-image-dpi", "150", "--min-keep-next", "4"], capture_output=True, text=True)
            self.assertEqual(docx_qa.returncode, 0, docx_qa.stdout + docx_qa.stderr)
            docx_report = json.loads(docx_qa.stdout)
            self.assertTrue(all(item["keep_next_count"] >= 4 for item in docx_report["documents"]))

    def test_qa_rejects_low_resolution_images(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            image = root / "low.png"
            Image.new("RGB", (100, 100), "white").save(image)
            pdf = root / "low.pdf"
            drawing = canvas.Canvas(str(pdf))
            drawing.drawImage(str(image), 72, 420, width=300, height=300)
            drawing.save()
            pdf_qa = subprocess.run(
                [sys.executable, str(SKILL / "scripts" / "qa_material.py"), str(pdf), "--min-images", "1", "--min-image-dpi", "150"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(pdf_qa.returncode, 0, pdf_qa.stdout + pdf_qa.stderr)
            self.assertTrue(json.loads(pdf_qa.stdout)["low_dpi_images"])

            docx = root / "low.docx"
            document = Document()
            document.add_picture(str(image), width=Mm(160))
            document.save(docx)
            docx_qa = subprocess.run(
                [sys.executable, str(SKILL / "scripts" / "qa_docx.py"), str(docx), "--min-image-dpi", "150"],
                capture_output=True,
                text=True,
            )
            self.assertNotEqual(docx_qa.returncode, 0, docx_qa.stdout + docx_qa.stderr)
            self.assertTrue(json.loads(docx_qa.stdout)["documents"][0]["low_dpi_images"])

    def test_legacy_print_kit_manifest_still_builds(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            manifest = root / "legacy.json"
            manifest.write_text(json.dumps({
                "mode": "memorization", "title": "Legacy", "subject": "Math", "chapter": "One",
                "print_profile": "bw", "sections": [{"title": "Section", "blocks": [{"type": "paragraph", "text": "Existing manifest content"}]}],
            }), encoding="utf-8")
            result = subprocess.run([sys.executable, str(ROOT / "skills" / "kaoyan-print-kit" / "scripts" / "build_material.py"), str(manifest), "--output-dir", str(root / "out")], capture_output=True, text=True)
            self.assertEqual(result.returncode, 0, result.stderr)
            self.assertEqual(len(list((root / "out").glob("*.pdf"))), 1)
            self.assertEqual(len(list((root / "out").glob("*.docx"))), 1)

    def test_shared_engine_keeps_mistakes_and_diagnostic_modes(self) -> None:
        fixtures = {
            "mistakes": {
                "type": "question", "id": "q1", "number": "1", "stem": "Compute 1 + 1.",
                "answer": "2", "analysis": "Direct addition.", "answer_space_lines": 2,
            },
            "diagnostic": {
                "type": "diagnostic", "id": "d1", "level": "recall", "prompt": "State the definition.",
                "answer": "A checked definition.", "rubric": "Include every condition.", "answer_space_lines": 2,
            },
        }
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            for mode, block in fixtures.items():
                manifest = root / f"{mode}.json"
                manifest.write_text(json.dumps({
                    "mode": mode,
                    "title": mode,
                    "subject": "Math",
                    "chapter": "One",
                    "print_profile": "bw",
                    "sections": [{"title": "Section", "blocks": [block]}],
                }), encoding="utf-8")
                output = root / mode
                result = subprocess.run(
                    [sys.executable, str(ROOT / "skills" / "kaoyan-print-kit" / "scripts" / "build_material.py"), str(manifest), "--output-dir", str(output)],
                    capture_output=True,
                    text=True,
                )
                self.assertEqual(result.returncode, 0, result.stderr)
                self.assertEqual(len(list(output.glob("*.pdf"))), 2)
                self.assertEqual(len(list(output.glob("*.docx"))), 2)

    def test_density_changes_docx_typography(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            sizes = {}
            for density in ("compact", "spacious"):
                manifest = root / f"{density}.json"
                manifest.write_text(json.dumps({
                    "mode": "memorization", "title": density, "subject": "Math", "chapter": "One",
                    "print_profile": "bw", "density": density,
                    "sections": [{"title": "Section", "blocks": [{"type": "paragraph", "text": "Content"}]}],
                }), encoding="utf-8")
                output = root / density
                result = subprocess.run([sys.executable, str(ROOT / "skills" / "kaoyan-print-kit" / "scripts" / "build_material.py"), str(manifest), "--output-dir", str(output)], capture_output=True, text=True)
                self.assertEqual(result.returncode, 0, result.stderr)
                document = Document(str(next(output.glob("*.docx"))))
                sizes[density] = document.styles["Normal"].font.size.pt
            self.assertLess(sizes["compact"], sizes["spacious"])

    def test_long_knowledge_segment_splits_across_pdf_pages(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            manifest = root / "long.json"
            manifest.write_text(json.dumps({
                "mode": "memorization", "title": "Long", "subject": "Math", "chapter": "One", "print_profile": "bw",
                "sections": [{"title": "Section", "blocks": [{
                    "type": "knowledge_point", "id": "long", "title": "Long source", "importance": "A", "personal_weak": False,
                    "grading_evidence": "user_designation: confirmed", "segments": [{"origin": "source_text", "origin_label": "来源原文", "text": "长段落内容。" * 5000, "source": "source.txt · 全文"}],
                }]}],
            }, ensure_ascii=False), encoding="utf-8")
            output = root / "out"
            result = subprocess.run([sys.executable, str(ROOT / "skills" / "kaoyan-print-kit" / "scripts" / "build_material.py"), str(manifest), "--output-dir", str(output)], capture_output=True, text=True)
            self.assertEqual(result.returncode, 0, result.stderr)
            pdf = next(output.glob("*.pdf"))
            self.assertGreater(len(PdfReader(str(pdf)).pages), 1)

    def test_knowledge_point_heading_stays_with_first_panel(self) -> None:
        with tempfile.TemporaryDirectory() as folder:
            root = Path(folder)
            blocks = []
            for index in range(1, 61):
                blocks.append({
                    "type": "knowledge_point", "id": f"p{index}", "title": f"Point-{index:02d}",
                    "importance": "B", "personal_weak": False, "grading_evidence": "user_designation: confirmed",
                    "segments": [{"origin": "source_text", "origin_label": "Source", "text": f"Body-{index:02d}", "source": "source.txt"}],
                })
            manifest = root / "many.json"
            manifest.write_text(json.dumps({
                "mode": "memorization", "title": "Many", "subject": "Math", "chapter": "One",
                "print_profile": "bw", "sections": [{"title": "Section", "blocks": blocks}],
            }), encoding="utf-8")
            output = root / "out"
            result = subprocess.run([sys.executable, str(ROOT / "skills" / "kaoyan-print-kit" / "scripts" / "build_material.py"), str(manifest), "--output-dir", str(output)], capture_output=True, text=True)
            self.assertEqual(result.returncode, 0, result.stderr)
            page_text = [page.extract_text() or "" for page in PdfReader(str(next(output.glob("*.pdf")))).pages]
            for index in range(1, 61):
                title = f"Point-{index:02d}"
                body = f"Body-{index:02d}"
                self.assertTrue(any(title in text and body in text for text in page_text), f"{title} was orphaned from its first panel")


if __name__ == "__main__":
    unittest.main()
