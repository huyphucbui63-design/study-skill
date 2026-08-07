#!/usr/bin/env python3
"""Audit DOCX structure and compare color/BW content without claiming visual QA."""

from __future__ import annotations

import argparse
from io import BytesIO
import json
from pathlib import Path
import sys
import zipfile

from docx import Document
from PIL import Image


def inspect(path: Path, min_image_dpi: float = 150.0, min_keep_next: int = 0) -> dict[str, object]:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_text = [[cell.text for cell in row.cells] for table in document.tables for row in table.rows]
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
        styles = archive.read("word/styles.xml").decode("utf-8")
    image_metrics = []
    unmeasured_images = 0
    for index, shape in enumerate(document.inline_shapes, 1):
        try:
            relationship_id = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = document.part.related_parts[relationship_id]
            with Image.open(BytesIO(image_part.blob)) as image:
                pixel_width, pixel_height = image.size
            width_inches = shape.width.inches
            height_inches = shape.height.inches
            if width_inches <= 0 or height_inches <= 0:
                raise ValueError("image has a non-positive display size")
            image_metrics.append({
                "index": index,
                "pixels": [pixel_width, pixel_height],
                "display_inches": [round(width_inches, 3), round(height_inches, 3)],
                "effective_dpi": round(min(pixel_width / width_inches, pixel_height / height_inches), 2),
            })
        except (AttributeError, KeyError, OSError, ValueError):
            unmeasured_images += 1
    low_dpi_images = [
        item for item in image_metrics
        if float(item["effective_dpi"]) + 0.05 < min_image_dpi
    ]
    keep_next_count = xml.count("<w:keepNext")
    return {
        "path": str(path), "paragraphs": paragraphs, "table_rows": table_text,
        "has_toc_field": " TOC " in xml, "has_east_asia_font": "w:eastAsia" in styles,
        "tables": len(document.tables), "images": len(document.inline_shapes),
        "image_metrics": image_metrics, "minimum_image_dpi": min_image_dpi,
        "low_dpi_images": low_dpi_images, "unmeasured_images": unmeasured_images,
        "keep_next_count": keep_next_count, "minimum_keep_next": min_keep_next,
        "ok": bool(paragraphs or table_text) and "w:eastAsia" in styles
        and not low_dpi_images and not unmeasured_images and keep_next_count >= min_keep_next,
    }


def content_signature(result: dict[str, object]) -> tuple[tuple[str, ...], tuple[tuple[str, ...], ...]]:
    return tuple(result["paragraphs"]), tuple(tuple(row) for row in result["table_rows"])


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("docx", nargs="+", type=Path)
    parser.add_argument("--require-toc", action="store_true")
    parser.add_argument("--min-image-dpi", type=float, default=150.0)
    parser.add_argument("--min-keep-next", type=int, default=0)
    args = parser.parse_args()
    try:
        if args.min_image_dpi < 0 or args.min_keep_next < 0:
            raise ValueError("minimum DPI and keep-next count cannot be negative")
        results = [
            inspect(path.expanduser().resolve(), args.min_image_dpi, args.min_keep_next)
            for path in args.docx
        ]
        equivalent = len({content_signature(result) for result in results}) == 1
        ok = all(result["ok"] for result in results) and equivalent
        if args.require_toc:
            ok = ok and all(result["has_toc_field"] for result in results)
        payload = {"documents": results, "content_equivalent": equivalent, "ok": ok}
        print(json.dumps(payload, ensure_ascii=False, indent=2))
        return 0 if ok else 1
    except (OSError, ValueError, KeyError, zipfile.BadZipFile) as exc:
        print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
