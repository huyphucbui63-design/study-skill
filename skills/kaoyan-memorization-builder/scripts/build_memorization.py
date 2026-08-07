#!/usr/bin/env python3
"""Validate a confirmed project and create color and black-and-white study editions."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
import shutil
import subprocess
import sys
import tempfile
from typing import Any

from docx import Document
from pypdf import PdfReader

from project_model import load_project, to_print_manifest


def resolve_media_paths(project: dict[str, Any], project_dir: Path) -> None:
    """Anchor project-relative visual assets before writing temporary manifests."""
    for chapter in project.get("chapters", []):
        for point in chapter.get("points", []):
            for visual in point.get("visuals", []):
                if visual.get("type") != "image" or not visual.get("path"):
                    continue
                path = Path(visual["path"]).expanduser()
                if not path.is_absolute():
                    path = project_dir / path
                path = path.resolve()
                if not path.is_file():
                    raise ValueError(f"Visual asset not found: {path}")
                visual["path"] = str(path)


def run_generator(manifest_path: Path, output_dir: Path, print_kit_script: Path) -> dict[str, Path]:
    command = [sys.executable, str(print_kit_script), str(manifest_path), "--output-dir", str(output_dir), "--variants", "study"]
    completed = subprocess.run(command, capture_output=True, text=True, check=False)
    if completed.returncode:
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "print-kit generation failed")
    payload = json.loads(completed.stdout)
    if not isinstance(payload, dict) or not isinstance(payload.get("generated"), list):
        raise RuntimeError("print-kit returned an invalid generated-file payload")
    generated = payload["generated"]
    if len(generated) != 2 or any(not isinstance(item, str) or not item.strip() for item in generated):
        raise RuntimeError("print-kit must return exactly one PDF and one DOCX")

    expected_root = output_dir.resolve()
    by_suffix: dict[str, Path] = {}
    for item in generated:
        candidate = Path(item).expanduser()
        if not candidate.is_absolute():
            candidate = output_dir / candidate
        candidate = candidate.resolve()
        try:
            candidate.relative_to(expected_root)
        except ValueError as exc:
            raise RuntimeError(f"print-kit returned a file outside its staging directory: {candidate}") from exc
        suffix = candidate.suffix.lower()
        if suffix not in {".pdf", ".docx"} or suffix in by_suffix:
            raise RuntimeError("print-kit must return exactly one PDF and one DOCX")
        if not candidate.is_file() or candidate.stat().st_size <= 0:
            raise RuntimeError(f"print-kit returned a missing or empty file: {candidate}")
        try:
            if suffix == ".pdf":
                if not PdfReader(str(candidate)).pages:
                    raise ValueError("PDF has no pages")
            else:
                document = Document(str(candidate))
                if not document.paragraphs and not document.tables and not document.inline_shapes:
                    raise ValueError("DOCX has no content")
        except Exception as exc:
            raise RuntimeError(f"print-kit returned an unreadable {suffix[1:].upper()} file: {candidate}") from exc
        by_suffix[suffix] = candidate
    if set(by_suffix) != {".pdf", ".docx"}:
        raise RuntimeError("print-kit must return exactly one PDF and one DOCX")
    return by_suffix


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("project", type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--print-kit-script", type=Path)
    args = parser.parse_args()
    skill_dir = Path(__file__).resolve().parents[2]
    default_script = skill_dir / "kaoyan-print-kit" / "scripts" / "build_material.py"
    print_kit_script = (args.print_kit_script or default_script).resolve()
    try:
        project_path = args.project.expanduser().resolve()
        project = load_project(project_path)
        resolve_media_paths(project, project_path.parent)
        output_dir = args.output_dir.expanduser().resolve()
        output_dir.parent.mkdir(parents=True, exist_ok=True)
        staged: dict[str, Path] = {}
        with tempfile.TemporaryDirectory(prefix="kaoyan-memo-") as temp:
            temp_dir = Path(temp)
            for profile in ("color", "bw"):
                manifest_path = temp_dir / f"manifest-{profile}.json"
                manifest_path.write_text(json.dumps(to_print_manifest(project, profile), ensure_ascii=False, indent=2), encoding="utf-8")
                profile_dir = temp_dir / profile
                files = run_generator(manifest_path, profile_dir, print_kit_script)
                for suffix, file_path in files.items():
                    destination = temp_dir / "publish" / f"{profile}-study{suffix}"
                    destination.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copy2(file_path, destination)
                    if destination.stat().st_size <= 0:
                        raise RuntimeError(f"Could not stage generated file: {destination}")
                    staged[destination.name] = destination

            expected_names = {
                "color-study.pdf", "color-study.docx", "bw-study.pdf", "bw-study.docx",
            }
            if set(staged) != expected_names:
                raise RuntimeError("Generation did not produce all four required study files")

            output_dir.mkdir(parents=True, exist_ok=True)
            generated = []
            for name in sorted(expected_names):
                destination = output_dir / name
                shutil.copy2(staged[name], destination)
                generated.append(str(destination))
        print(json.dumps({"generated": generated, "profiles": ["color", "bw"]}, ensure_ascii=False, indent=2))
        return 0
    except (OSError, ValueError, RuntimeError, json.JSONDecodeError) as exc:
        print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
