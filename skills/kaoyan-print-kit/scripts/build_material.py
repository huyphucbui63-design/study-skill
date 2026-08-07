#!/usr/bin/env python3
"""Build print-ready study PDF and DOCX files from a UTF-8 JSON manifest."""

from __future__ import annotations

import argparse
import html
import json
import os
import re
import sys
from datetime import date
from pathlib import Path
from typing import Any

from PIL import Image as PILImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Flowable,
    Frame,
    Image,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents


MODES = {"memorization", "mistakes", "diagnostic"}
PRINT_PROFILES = {"bw", "color"}
DENSITIES = {"compact", "standard", "spacious"}
IMAGE_TARGET_DPI = 150.0
PDF_FONT_CANDIDATES = (
    (Path("C:/Windows/Fonts/Deng.ttf"), Path("C:/Windows/Fonts/Dengb.ttf")),
    (Path("C:/Windows/Fonts/NotoSansSC-Regular.ttf"), Path("C:/Windows/Fonts/NotoSansSC-Bold.ttf")),
    (Path("/usr/share/fonts/truetype/noto/NotoSansCJKsc-Regular.ttf"), Path("/usr/share/fonts/truetype/noto/NotoSansCJKsc-Bold.ttf")),
    (Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"), Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc")),
    (Path("/System/Library/Fonts/PingFang.ttc"), Path("/System/Library/Fonts/PingFang.ttc")),
)
DOCX_FONT = os.environ.get(
    "KAOYAN_DOCX_FONT",
    "DengXian" if sys.platform == "win32" else "PingFang SC" if sys.platform == "darwin" else "Noto Sans CJK SC",
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("manifest", type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument(
        "--variants",
        help="Comma-separated variants: study,practice,answers",
    )
    return parser.parse_args()


def load_manifest(path: Path) -> dict[str, Any]:
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"Could not read manifest: {exc}") from exc

    required = ("mode", "title", "subject", "chapter", "print_profile", "sections")
    missing = [key for key in required if not data.get(key)]
    if missing:
        raise ValueError(f"Missing required manifest fields: {', '.join(missing)}")
    if data["mode"] not in MODES:
        raise ValueError(f"Unsupported mode: {data['mode']}")
    if data["print_profile"] not in PRINT_PROFILES:
        raise ValueError("print_profile must be bw or color")
    if data.get("density", "standard") not in DENSITIES:
        raise ValueError("density must be compact, standard, or spacious")
    if not isinstance(data["sections"], list):
        raise ValueError("sections must be a list")
    return data


def variants_for(manifest: dict[str, Any], requested: str | None) -> list[str]:
    if requested:
        variants = [item.strip() for item in requested.split(",") if item.strip()]
    elif manifest["mode"] == "memorization":
        variants = ["study"]
    else:
        variants = ["practice", "answers"]
    invalid = sorted(set(variants) - {"study", "practice", "answers"})
    if invalid:
        raise ValueError(f"Unsupported variants: {', '.join(invalid)}")
    return variants


def safe_name(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", value).strip(" .")
    return re.sub(r"\s+", "_", cleaned) or "study-material"


def resolve_media(path: str, manifest_dir: Path) -> Path:
    candidate = Path(path).expanduser()
    if not candidate.is_absolute():
        candidate = manifest_dir / candidate
    candidate = candidate.resolve()
    if not candidate.is_file():
        raise ValueError(f"Media file not found: {candidate}")
    return candidate


def register_pdf_fonts() -> tuple[str, str]:
    configured_regular = os.environ.get("KAOYAN_FONT_REGULAR")
    configured_bold = os.environ.get("KAOYAN_FONT_BOLD")
    if bool(configured_regular) != bool(configured_bold):
        raise RuntimeError("Set both KAOYAN_FONT_REGULAR and KAOYAN_FONT_BOLD, or neither.")

    candidates = list(PDF_FONT_CANDIDATES)
    if configured_regular and configured_bold:
        candidates.insert(0, (Path(configured_regular), Path(configured_bold)))

    errors: list[str] = []
    for regular_path, bold_path in candidates:
        if not regular_path.is_file() or not bold_path.is_file():
            continue
        try:
            pdfmetrics.registerFont(TTFont("KaoyanCJK", str(regular_path)))
            pdfmetrics.registerFont(TTFont("KaoyanCJK-Bold", str(bold_path)))
            return "KaoyanCJK", "KaoyanCJK-Bold"
        except Exception as exc:  # ReportLab support varies for variable fonts and TTC files.
            errors.append(f"{regular_path}: {exc}")

    detail = f" Tried: {'; '.join(errors)}" if errors else ""
    raise RuntimeError(
        "No usable Simplified Chinese TrueType font was found. Install Noto Sans CJK SC "
        "or set KAOYAN_FONT_REGULAR and KAOYAN_FONT_BOLD to font files." + detail
    )


def rich_text(value: Any) -> str:
    return html.escape(str(value or "")).replace("\n", "<br/>")


class RuledSpace(Flowable):
    def __init__(self, lines: int, width: float, color: colors.Color):
        super().__init__()
        self.lines = max(1, min(int(lines), 20))
        self.width = width
        self.line_height = 8 * mm
        self.height = self.lines * self.line_height
        self.color = color

    def draw(self) -> None:
        self.canv.setStrokeColor(self.color)
        self.canv.setLineWidth(0.45)
        for index in range(self.lines):
            y = self.height - (index + 1) * self.line_height + 2 * mm
            self.canv.line(0, y, self.width, y)


def pdf_palette(profile: str) -> dict[str, colors.Color]:
    if profile == "bw":
        return {
            "primary": colors.HexColor("#202020"),
            "secondary": colors.HexColor("#555555"),
            "accent": colors.HexColor("#707070"),
            "light": colors.HexColor("#F0F0F0"),
            "rule": colors.HexColor("#A0A0A0"),
            "warning": colors.HexColor("#303030"),
        }
    return {
        "primary": colors.HexColor("#17324D"),
        "secondary": colors.HexColor("#35606F"),
        "accent": colors.HexColor("#167C80"),
        "light": colors.HexColor("#EAF3F2"),
        "rule": colors.HexColor("#8CA5AA"),
        "warning": colors.HexColor("#A33636"),
    }


def pdf_styles(profile: str, regular: str, bold: str, density: str = "standard") -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    palette = pdf_palette(profile)
    density_tokens = {
        "compact": {"body": 9.5, "leading": 13.2, "after": 1.8, "small": 7.5, "label": 10, "section": 12.5},
        "standard": {"body": 10.5, "leading": 15.2, "after": 2.5, "small": 8, "label": 10.5, "section": 13},
        "spacious": {"body": 11, "leading": 16.8, "after": 3.2, "small": 8.5, "label": 11, "section": 13.5},
    }[density]
    return {
        "title": ParagraphStyle(
            "StudyTitle", parent=base["Title"], fontName=bold, fontSize=16,
            leading=21, alignment=TA_CENTER, textColor=palette["primary"],
            spaceAfter=7 * mm,
        ),
        "meta": ParagraphStyle(
            "StudyMeta", parent=base["Normal"], fontName=regular, fontSize=9,
            leading=13, alignment=TA_CENTER, textColor=palette["secondary"],
            spaceAfter=5 * mm,
        ),
        "section": ParagraphStyle(
            "StudySection", parent=base["Heading2"], fontName=bold, fontSize=density_tokens["section"],
            leading=18, textColor=palette["primary"], spaceBefore=4 * mm,
            spaceAfter=2.5 * mm, keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "StudyBody", parent=base["BodyText"], fontName=regular, fontSize=density_tokens["body"],
            leading=density_tokens["leading"], textColor=colors.black, spaceAfter=density_tokens["after"] * mm,
            wordWrap="CJK",
        ),
        "small": ParagraphStyle(
            "StudySmall", parent=base["BodyText"], fontName=regular, fontSize=density_tokens["small"],
            leading=11, textColor=palette["secondary"], spaceAfter=1.5 * mm,
            wordWrap="CJK",
        ),
        "label": ParagraphStyle(
            "StudyLabel", parent=base["BodyText"], fontName=bold, fontSize=density_tokens["label"],
            leading=15, textColor=palette["primary"], spaceAfter=1.5 * mm,
            wordWrap="CJK",
        ),
        "toc": ParagraphStyle(
            "StudyToc", parent=base["BodyText"], fontName=regular, fontSize=10,
            leading=15, textColor=palette["primary"], leftIndent=4 * mm,
            firstLineIndent=-4 * mm, wordWrap="CJK",
        ),
        "origin": ParagraphStyle(
            "StudyOrigin", parent=base["BodyText"], fontName=bold, fontSize=density_tokens["small"] + 0.5,
            leading=12, textColor=palette["secondary"], spaceAfter=1 * mm,
            wordWrap="CJK",
        ),
    }


class KaoyanDocTemplate(BaseDocTemplate):
    """Collect section headings for a deterministic PDF table of contents."""

    def afterFlowable(self, flowable: Any) -> None:
        bookmark = getattr(flowable, "_bookmarkName", None)
        if not bookmark:
            return
        title = flowable.getPlainText()
        self.canv.bookmarkPage(bookmark)
        self.canv.addOutlineEntry(title, bookmark, level=0, closed=False)
        self.notify("TOCEntry", (0, title, self.page, bookmark))


def image_display_size(
    path: Path,
    max_width: float,
    max_height: float,
    target_dpi: float = IMAGE_TARGET_DPI,
) -> tuple[float, float]:
    with PILImage.open(path) as image:
        width, height = image.size
    if width <= 0 or height <= 0:
        raise ValueError(f"Image has invalid pixel dimensions: {path}")
    width_points = width * 72.0 / target_dpi
    height_points = height * 72.0 / target_dpi
    scale = min(max_width / width_points, max_height / height_points, 1.0)
    return width_points * scale, height_points * scale


def pdf_image(path: Path, max_width: float = 160 * mm, max_height: float = 170 * mm) -> Image:
    width, height = image_display_size(path, max_width, max_height)
    return Image(str(path), width=width, height=height)


def add_docx_picture(
    document: Document,
    path: Path,
    max_width: float = 160 * mm,
    max_height: float = 170 * mm,
) -> None:
    width, height = image_display_size(path, max_width, max_height)
    document.add_picture(str(path), width=Pt(width), height=Pt(height))


def source_paragraph(block: dict[str, Any], styles: dict[str, ParagraphStyle]) -> list[Any]:
    source = block.get("source")
    return [Paragraph(f"来源：{rich_text(source)}", styles["small"])] if source else []


def answer_panel(label: str, text: Any, styles: dict[str, ParagraphStyle], palette: dict[str, colors.Color]) -> Table:
    table = Table([[Paragraph(f"<b>{rich_text(label)}</b><br/>{rich_text(text)}", styles["body"])]], colWidths=[166 * mm])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), palette["light"]),
        ("BOX", (0, 0), (-1, -1), 0.6, palette["rule"]),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    return table


def memory_structure_pdf(
    block: dict[str, Any], styles: dict[str, ParagraphStyle], palette: dict[str, colors.Color]
) -> list[Any]:
    kind = block["type"]
    title = Paragraph(rich_text(block.get("title", "")), styles["label"])
    prefix: list[Any] = [title]
    if block.get("origin_label"):
        prefix.append(Paragraph(rich_text(block["origin_label"]), styles["origin"]))
    if block.get("source"):
        prefix.append(Paragraph(f"来源：{rich_text(block['source'])}", styles["small"]))
    if kind == "comparison":
        headers = block.get("headers", [])
        rows = block.get("rows", [])
        if not headers or not rows or any(len(row) != len(headers) for row in rows):
            raise ValueError("comparison requires equal-width headers and rows")
        width = 166 * mm / len(headers)
        data = [[Paragraph(rich_text(cell), styles["origin"]) for cell in headers]]
        data.extend([Paragraph(rich_text(cell), styles["body"]) for cell in row] for row in rows)
        table = Table(data, colWidths=[width] * len(headers), repeatRows=1)
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.7, palette["rule"]),
            ("INNERGRID", (0, 0), (-1, -1), 0.35, palette["rule"]),
            ("BACKGROUND", (0, 0), (-1, 0), palette["light"]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return [*prefix, table, Spacer(1, 2.5 * mm)]
    if kind in {"process", "timeline"}:
        items = block.get("items", [])
        if not items:
            raise ValueError(f"{kind} requires items")
        rows = []
        for index, item in enumerate(items, 1):
            marker = str(index) if kind == "process" else "T" + str(index)
            rows.append([Paragraph(f"<b>{marker}</b>", styles["label"]), Paragraph(rich_text(item), styles["body"])])
        table = Table(rows, colWidths=[16 * mm, 150 * mm])
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.6, palette["rule"]),
            ("LINEBELOW", (0, 0), (-1, -2), 0.35, palette["rule"]),
            ("BACKGROUND", (0, 0), (0, -1), palette["light"]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return [*prefix, KeepTogether(table), Spacer(1, 2.5 * mm)]
    relations = block.get("relations", [])
    if not relations:
        raise ValueError("relationship requires relations")
    rows = [[
        Paragraph(rich_text(item.get("from")), styles["body"]),
        Paragraph(rich_text(item.get("relation")), styles["origin"]),
        Paragraph(rich_text(item.get("to")), styles["body"]),
    ] for item in relations]
    table = Table(rows, colWidths=[66 * mm, 34 * mm, 66 * mm])
    table.setStyle(TableStyle([
        ("BOX", (0, 0), (-1, -1), 0.6, palette["rule"]),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, palette["rule"]),
        ("BACKGROUND", (1, 0), (1, -1), palette["light"]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return [*prefix, table, Spacer(1, 2.5 * mm)]


def block_to_pdf(
    block: dict[str, Any],
    variant: str,
    styles: dict[str, ParagraphStyle],
    palette: dict[str, colors.Color],
    manifest_dir: Path,
) -> list[Any]:
    kind = block.get("type", "paragraph")
    items: list[Any] = []
    if kind == "paragraph":
        items.append(Paragraph(rich_text(block.get("text")), styles["body"]))
    elif kind == "bullets":
        for item in block.get("items", []):
            items.append(Paragraph(f"• {rich_text(item)}", styles["body"]))
    elif kind == "callout":
        items.append(answer_panel(block.get("label", "提示"), block.get("text"), styles, palette))
        items.append(Spacer(1, 2.5 * mm))
    elif kind == "formula":
        group = [Paragraph(rich_text(block.get("formula")), styles["label"])]
        if block.get("notes"):
            group.append(Paragraph(rich_text(block["notes"]), styles["body"]))
        group.extend(source_paragraph(block, styles))
        items.append(KeepTogether(group))
    elif kind == "definition":
        rows = [
            [Paragraph("定义", styles["label"]), Paragraph(rich_text(block.get("definition")), styles["body"])],
        ]
        if block.get("keywords"):
            rows.append([Paragraph("关键词", styles["label"]), Paragraph("、".join(rich_text(x) for x in block["keywords"]), styles["body"])])
        for key, label in (("boundary", "边界"), ("counterexample", "反例")):
            if block.get(key):
                rows.append([Paragraph(label, styles["label"]), Paragraph(rich_text(block[key]), styles["body"])])
        table = Table(rows, colWidths=[25 * mm, 141 * mm], repeatRows=0)
        table.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.6, palette["rule"]),
            ("INNERGRID", (0, 0), (-1, -1), 0.35, palette["rule"]),
            ("BACKGROUND", (0, 0), (0, -1), palette["light"]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        if block.get("term"):
            items.append(Paragraph(rich_text(block["term"]), styles["label"]))
        items.extend([table, Spacer(1, 2.5 * mm), *source_paragraph(block, styles)])
    elif kind == "knowledge_point":
        importance = block.get("importance")
        if importance not in {"A", "B", "C"}:
            raise ValueError("knowledge_point importance must be A, B, or C")
        markers = f"[{importance}]" + (" [R]" if block.get("personal_weak") else "")
        heading = Table([[
            Paragraph(markers, styles["label"]),
            Paragraph(rich_text(block.get("title")), styles["label"]),
        ]], colWidths=[28 * mm, 138 * mm])
        heading.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.8, palette["primary"]),
            ("BACKGROUND", (0, 0), (0, 0), palette["light"]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN", (0, 0), (0, 0), "CENTER"),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        heading.keepWithNext = True
        heading_gap = Spacer(1, 1.5 * mm)
        heading_gap.keepWithNext = True
        items.extend([heading, heading_gap])
        for segment in block.get("segments", []):
            label = segment.get("origin_label") or segment.get("origin")
            source_text = ""
            if segment.get("source"):
                source_text = f'<br/><font size="8" color="{palette["secondary"].hexval()}">来源：{rich_text(segment["source"])}</font>'
            fill = colors.white if segment.get("origin") in {"source_text", "user_note"} else palette["light"]
            panel_style = ParagraphStyle(
                f"KnowledgePanel-{segment.get('origin', 'unknown')}", parent=styles["body"],
                backColor=fill, borderColor=palette["accent"], borderWidth=0.7,
                borderPadding=(6, 7, 6, 8), spaceAfter=1.5 * mm,
                leftIndent=7 * mm, rightIndent=7 * mm, splitLongWords=True,
            )
            panel = Paragraph(
                f"<b>{rich_text(label)}</b><br/>{rich_text(segment.get('text'))}{source_text}",
                panel_style,
            )
            items.extend([panel, Spacer(1, 1.5 * mm)])
        if block.get("grading_evidence"):
            items.append(Paragraph(f"分级依据：{rich_text(block['grading_evidence'])}", styles["small"]))
        items.append(Spacer(1, 2.5 * mm))
    elif kind in {"comparison", "process", "timeline", "relationship"}:
        items.extend(memory_structure_pdf(block, styles, palette))
    elif kind == "recall":
        items.append(Paragraph(rich_text(block.get("label", "主动回忆")), styles["origin"]))
        items.extend([RuledSpace(block.get("lines", 3), 166 * mm, palette["rule"]), Spacer(1, 3 * mm)])
    elif kind == "image":
        image_path = resolve_media(block["path"], manifest_dir)
        group = [pdf_image(image_path)]
        if block.get("caption"):
            group.append(Paragraph(rich_text(block["caption"]), styles["small"]))
        group.extend(source_paragraph(block, styles))
        items.append(KeepTogether(group))
    elif kind in {"question", "diagnostic"}:
        if kind == "question":
            heading = f"{block.get('number') or block.get('id') or '题目'}  {block.get('stem', '')}"
        else:
            heading = f"[{block.get('level', '理解检测')}] {block.get('prompt', '')}"
        group: list[Any] = [Paragraph(rich_text(heading), styles["label"])]
        for option in block.get("options", []):
            if isinstance(option, dict):
                option = f"{option.get('label', '')}. {option.get('text', '')}"
            group.append(Paragraph(rich_text(option), styles["body"]))
        for subquestion in block.get("subquestions", []):
            if isinstance(subquestion, dict):
                subquestion = f"{subquestion.get('number', '')} {subquestion.get('text', '')}"
            group.append(Paragraph(rich_text(subquestion), styles["body"]))
        for table_data in block.get("tables", []):
            raw_rows = ([table_data.get("headers", [])] if table_data.get("headers") else []) + table_data.get("rows", [])
            if not raw_rows:
                continue
            column_count = max(len(row) for row in raw_rows)
            rows = [
                [Paragraph(rich_text(row[index] if index < len(row) else ""), styles["body"]) for index in range(column_count)]
                for row in raw_rows
            ]
            table = Table(rows, colWidths=[155 * mm / column_count] * column_count, repeatRows=1 if table_data.get("headers") else 0)
            commands = [
                ("BOX", (0, 0), (-1, -1), 0.6, palette["rule"]),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, palette["rule"]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
            if table_data.get("headers"):
                commands.append(("BACKGROUND", (0, 0), (-1, 0), palette["light"]))
            table.setStyle(TableStyle(commands))
            group.extend([table, Spacer(1, 2 * mm)])
        image_value = block.get("image")
        if image_value:
            group.append(pdf_image(resolve_media(image_value, manifest_dir), 155 * mm, 120 * mm))
        for graphic in block.get("graphics", []):
            value = graphic.get("path") if isinstance(graphic, dict) else graphic
            if value:
                group.append(pdf_image(resolve_media(value, manifest_dir), 155 * mm, 120 * mm))
        group.extend(source_paragraph(block, styles))
        if variant == "practice":
            items.append(KeepTogether([
                *group,
                RuledSpace(block.get("answer_space_lines", 5), 166 * mm, palette["rule"]),
                Spacer(1, 3 * mm),
            ]))
        else:
            items.append(KeepTogether(group))
            for key, label in (
                ("answer", "答案"), ("analysis", "解析"), ("error_cause", "错误原因"),
                ("knowledge_points", "关联知识点"), ("rubric", "评分标准"),
                ("misconceptions", "常见误解"),
            ):
                value = block.get(key)
                if isinstance(value, list):
                    value = "；".join(str(item) for item in value)
                if value:
                    items.extend([answer_panel(label, value, styles, palette), Spacer(1, 2 * mm)])
    else:
        raise ValueError(f"Unsupported block type: {kind}")
    return items


def build_pdf(manifest: dict[str, Any], variant: str, output: Path, manifest_dir: Path) -> None:
    regular, bold = register_pdf_fonts()
    styles = pdf_styles(manifest["print_profile"], regular, bold, manifest.get("density", "standard"))
    palette = pdf_palette(manifest["print_profile"])
    page_width, page_height = A4
    inner, outer, top, bottom = 18 * mm, 12 * mm, 14 * mm, 15 * mm

    def decorate(canvas: Any, doc: Any) -> None:
        canvas.saveState()
        canvas.setFont(regular, 8)
        canvas.setFillColor(palette["secondary"])
        canvas.drawString(outer, 8 * mm, safe_name(manifest["subject"]))
        canvas.drawCentredString(page_width / 2, 8 * mm, str(doc.page))
        canvas.restoreState()

    odd_frame = Frame(inner, bottom, page_width - inner - outer, page_height - top - bottom, id="odd")
    even_frame = Frame(outer, bottom, page_width - inner - outer, page_height - top - bottom, id="even")
    odd = PageTemplate(id="Odd", frames=[odd_frame], onPage=decorate, autoNextPageTemplate="Even")
    even = PageTemplate(id="Even", frames=[even_frame], onPage=decorate, autoNextPageTemplate="Odd")
    doc = KaoyanDocTemplate(str(output), pagesize=A4, pageTemplates=[odd, even], title=manifest["title"], author="Kaoyan Print Kit")

    variant_labels = {"study": "背诵版", "practice": "练习版", "answers": "解析版"}
    story: list[Any] = [
        Paragraph(rich_text(manifest["title"]), styles["title"]),
        Paragraph(
            f"{rich_text(manifest['subject'])} · {rich_text(manifest['chapter'])} · {variant_labels[variant]}",
            styles["meta"],
        ),
    ]
    if manifest.get("include_toc"):
        story.append(Paragraph("目录", styles["section"]))
        toc = TableOfContents()
        toc.levelStyles = [styles["toc"]]
        story.extend([toc, PageBreak()])
    for section_index, section in enumerate(manifest["sections"]):
        if section.get("page_break_before"):
            story.append(PageBreak())
        heading = Paragraph(rich_text(section.get("title", "")), styles["section"])
        heading._bookmarkName = f"section-{section.get('id') or section_index + 1}"
        story.append(heading)
        for block in section.get("blocks", []):
            story.extend(block_to_pdf(block, variant, styles, palette, manifest_dir))
    doc.multiBuild(story)
    if len(PdfReader(str(output)).pages) < 1:
        raise RuntimeError(f"Generated PDF has no pages: {output}")


def set_docx_font(run: Any, name: str = DOCX_FONT) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_table_geometry(table: Any, widths_mm: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    total_dxa = int(sum(widths_mm) * 56.692913)
    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.first_child_found_in("w:tblW")
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.append(tbl_width)
    tbl_width.set(qn("w:w"), str(total_dxa))
    tbl_width.set(qn("w:type"), "dxa")
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    tbl_pr.append(indent)
    for row in table.rows:
        for cell, width_mm in zip(row.cells, widths_mm):
            width_dxa = int(width_mm * 56.692913)
            cell.width = Mm(width_mm)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.first_child_found_in("w:tcW")
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(width_dxa))
            tc_width.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph: Any, color: str = "A0A0A0") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def add_docx_text(document: Document, text: Any, style: str | None = None, bold: bool = False) -> Any:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    set_docx_font(run)
    return paragraph


def add_page_field(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_toc_field(paragraph: Any) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-1" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开文档后更新目录域以显示页码。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.extend([begin, instruction, separate, placeholder, end])
    paragraph._p.append(run)


def configure_docx(document: Document, manifest: dict[str, Any]) -> None:
    section = document.sections[0]
    section.page_width, section.page_height = Mm(210), Mm(297)
    section.left_margin, section.right_margin = Mm(18), Mm(12)
    section.top_margin, section.bottom_margin = Mm(14), Mm(15)
    mirror = OxmlElement("w:mirrorMargins")
    section._sectPr.append(mirror)

    density = {
        "compact": {"body": 9.5, "after": 3, "line": 1.25, "title": 15, "h1": 12.5, "h2": 10},
        "standard": {"body": 10.5, "after": 5, "line": 1.45, "title": 16, "h1": 13, "h2": 11},
        "spacious": {"body": 11, "after": 7, "line": 1.6, "title": 17, "h1": 13.5, "h2": 11.5},
    }[manifest.get("density", "standard")]
    normal = document.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal.font.size = Pt(density["body"])
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
    normal.paragraph_format.space_after = Pt(density["after"])
    normal.paragraph_format.line_spacing = density["line"]
    for style_name, size, color, bold in (
        ("Title", density["title"], "17324D", True),
        ("Heading 1", density["h1"], "17324D", True),
        ("Heading 2", density["h2"], "35606F", True),
    ):
        style = document.styles[style_name]
        style.font.name = DOCX_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string("202020" if manifest["print_profile"] == "bw" else color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)

    list_style = document.styles["List Bullet"]
    list_style.font.name = DOCX_FONT
    list_style.font.size = Pt(density["body"])
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT)
    list_style.paragraph_format.left_indent = Mm(9.5)
    list_style.paragraph_format.first_line_indent = Mm(-4.75)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    add_page_field(footer)


def add_source_docx(document: Document, block: dict[str, Any]) -> None:
    if block.get("source"):
        paragraph = add_docx_text(document, f"来源：{block['source']}")
        for run in paragraph.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)


def add_docx_block(document: Document, block: dict[str, Any], variant: str, manifest_dir: Path, profile: str) -> None:
    kind = block.get("type", "paragraph")
    light_fill = "F0F0F0" if profile == "bw" else "EAF3F2"
    if kind == "paragraph":
        add_docx_text(document, block.get("text"))
    elif kind == "bullets":
        for item in block.get("items", []):
            add_docx_text(document, item, style="List Bullet")
    elif kind == "callout":
        table = document.add_table(rows=1, cols=1)
        set_table_geometry(table, [166])
        set_cell_shading(table.cell(0, 0), light_fill)
        paragraph = table.cell(0, 0).paragraphs[0]
        run = paragraph.add_run(f"{block.get('label', '提示')}\n{block.get('text', '')}")
        set_docx_font(run)
    elif kind == "formula":
        add_docx_text(document, block.get("formula"), style="Heading 2")
        if block.get("notes"):
            add_docx_text(document, block["notes"])
        add_source_docx(document, block)
    elif kind == "definition":
        if block.get("term"):
            add_docx_text(document, block["term"], style="Heading 2")
        table = document.add_table(rows=0, cols=2)
        rows = [("定义", block.get("definition", ""))]
        if block.get("keywords"):
            rows.append(("关键词", "、".join(str(x) for x in block["keywords"])))
        rows.extend((label, block[key]) for key, label in (("boundary", "边界"), ("counterexample", "反例")) if block.get(key))
        for label, value in rows:
            cells = table.add_row().cells
            set_cell_shading(cells[0], light_fill)
            for cell, text in zip(cells, (label, value)):
                run = cell.paragraphs[0].add_run(str(text))
                set_docx_font(run)
        set_table_geometry(table, [25, 141])
        add_source_docx(document, block)
    elif kind == "knowledge_point":
        importance = block.get("importance")
        if importance not in {"A", "B", "C"}:
            raise ValueError("knowledge_point importance must be A, B, or C")
        heading = document.add_table(rows=1, cols=2)
        heading.style = "Table Grid"
        set_table_geometry(heading, [28, 138])
        set_cell_shading(heading.cell(0, 0), light_fill)
        marker = f"[{importance}]" + (" [R]" if block.get("personal_weak") else "")
        for cell, value in zip(heading.rows[0].cells, (marker, block.get("title", ""))):
            run = cell.paragraphs[0].add_run(str(value))
            run.bold = True
            set_docx_font(run)
            cell.paragraphs[0].paragraph_format.keep_with_next = True
        heading.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for segment in block.get("segments", []):
            panel = document.add_table(rows=1, cols=1)
            panel.style = "Table Grid"
            set_table_geometry(panel, [166])
            if segment.get("origin") not in {"source_text", "user_note"}:
                set_cell_shading(panel.cell(0, 0), light_fill)
            paragraph = panel.cell(0, 0).paragraphs[0]
            label_run = paragraph.add_run(str(segment.get("origin_label") or segment.get("origin")))
            label_run.bold = True
            label_run.font.size = Pt(8.5)
            set_docx_font(label_run)
            paragraph.add_run("\n")
            content_run = paragraph.add_run(str(segment.get("text", "")))
            set_docx_font(content_run)
            if segment.get("source"):
                source_run = paragraph.add_run(f"\n来源：{segment['source']}")
                source_run.font.size = Pt(8)
                source_run.font.color.rgb = RGBColor(100, 100, 100)
                set_docx_font(source_run)
        if block.get("grading_evidence"):
            paragraph = add_docx_text(document, f"分级依据：{block['grading_evidence']}")
            for run in paragraph.runs:
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(100, 100, 100)
    elif kind in {"comparison", "process", "timeline", "relationship"}:
        add_docx_text(document, block.get("title", ""), style="Heading 2")
        if block.get("origin_label"):
            paragraph = add_docx_text(document, block["origin_label"], bold=True)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
        if block.get("source"):
            add_source_docx(document, block)
        if kind == "comparison":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if not headers or not rows or any(len(row) != len(headers) for row in rows):
                raise ValueError("comparison requires equal-width headers and rows")
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for cell, value in zip(table.rows[0].cells, headers):
                set_cell_shading(cell, light_fill)
                run = cell.paragraphs[0].add_run(str(value))
                run.bold = True
                set_docx_font(run)
            for row in rows:
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    run = cell.paragraphs[0].add_run(str(value))
                    set_docx_font(run)
            set_table_geometry(table, [166 / len(headers)] * len(headers))
        elif kind in {"process", "timeline"}:
            items = block.get("items", [])
            if not items:
                raise ValueError(f"{kind} requires items")
            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"
            for index, value in enumerate(items, 1):
                cells = table.add_row().cells
                set_cell_shading(cells[0], light_fill)
                marker = str(index) if kind == "process" else f"T{index}"
                for cell, text in zip(cells, (marker, value)):
                    run = cell.paragraphs[0].add_run(str(text))
                    set_docx_font(run)
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_table_geometry(table, [16, 150])
        else:
            relations = block.get("relations", [])
            if not relations:
                raise ValueError("relationship requires relations")
            table = document.add_table(rows=0, cols=3)
            table.style = "Table Grid"
            for relation in relations:
                cells = table.add_row().cells
                set_cell_shading(cells[1], light_fill)
                values = (relation.get("from", ""), relation.get("relation", ""), relation.get("to", ""))
                for cell, value in zip(cells, values):
                    run = cell.paragraphs[0].add_run(str(value))
                    set_docx_font(run)
                cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_table_geometry(table, [66, 34, 66])
    elif kind == "recall":
        paragraph = add_docx_text(document, block.get("label", "主动回忆"), bold=True)
        paragraph.paragraph_format.keep_with_next = True
        for _ in range(max(1, min(int(block.get("lines", 3)), 12))):
            line = document.add_paragraph()
            line.paragraph_format.space_after = Pt(7)
            add_bottom_border(line)
    elif kind == "image":
        image_path = resolve_media(block["path"], manifest_dir)
        add_docx_picture(document, image_path)
        if block.get("caption"):
            paragraph = add_docx_text(document, block["caption"])
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_source_docx(document, block)
    elif kind in {"question", "diagnostic"}:
        if kind == "question":
            label = f"{block.get('number') or block.get('id') or '题目'}  {block.get('stem', '')}"
        else:
            label = f"[{block.get('level', '理解检测')}] {block.get('prompt', '')}"
        add_docx_text(document, label, style="Heading 2")
        for option in block.get("options", []):
            if isinstance(option, dict):
                option = f"{option.get('label', '')}. {option.get('text', '')}"
            add_docx_text(document, option)
        for subquestion in block.get("subquestions", []):
            if isinstance(subquestion, dict):
                subquestion = f"{subquestion.get('number', '')} {subquestion.get('text', '')}"
            add_docx_text(document, subquestion)
        for table_data in block.get("tables", []):
            raw_rows = ([table_data.get("headers", [])] if table_data.get("headers") else []) + table_data.get("rows", [])
            if not raw_rows:
                continue
            column_count = max(len(row) for row in raw_rows)
            table = document.add_table(rows=0, cols=column_count)
            for row_index, values in enumerate(raw_rows):
                cells = table.add_row().cells
                for index, cell in enumerate(cells):
                    run = cell.paragraphs[0].add_run(str(values[index] if index < len(values) else ""))
                    set_docx_font(run)
                    if row_index == 0 and table_data.get("headers"):
                        run.bold = True
                        set_cell_shading(cell, light_fill)
            set_table_geometry(table, [155 / column_count] * column_count)
        if block.get("image"):
            add_docx_picture(
                document,
                resolve_media(block["image"], manifest_dir),
                155 * mm,
                120 * mm,
            )
        for graphic in block.get("graphics", []):
            value = graphic.get("path") if isinstance(graphic, dict) else graphic
            if value:
                add_docx_picture(document, resolve_media(value, manifest_dir), 155 * mm, 120 * mm)
        add_source_docx(document, block)
        if variant == "practice":
            for _ in range(max(1, min(int(block.get("answer_space_lines", 5)), 20))):
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(7)
                add_bottom_border(paragraph)
        else:
            for key, label_name in (
                ("answer", "答案"), ("analysis", "解析"), ("error_cause", "错误原因"),
                ("knowledge_points", "关联知识点"), ("rubric", "评分标准"),
                ("misconceptions", "常见误解"),
            ):
                value = block.get(key)
                if isinstance(value, list):
                    value = "；".join(str(item) for item in value)
                if value:
                    table = document.add_table(rows=1, cols=1)
                    set_table_geometry(table, [166])
                    set_cell_shading(table.cell(0, 0), light_fill)
                    run = table.cell(0, 0).paragraphs[0].add_run(f"{label_name}\n{value}")
                    set_docx_font(run)
    else:
        raise ValueError(f"Unsupported block type: {kind}")


def build_docx(manifest: dict[str, Any], variant: str, output: Path, manifest_dir: Path) -> None:
    document = Document()
    configure_docx(document, manifest)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(manifest["title"])
    set_docx_font(run)
    variant_labels = {"study": "背诵版", "practice": "练习版", "answers": "解析版"}
    meta = add_docx_text(document, f"{manifest['subject']} · {manifest['chapter']} · {variant_labels[variant]}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if manifest.get("include_toc"):
        add_docx_text(document, "目录", style="Heading 1")
        toc = document.add_paragraph()
        add_toc_field(toc)
        document.add_page_break()
    for section_index, section in enumerate(manifest["sections"]):
        if section.get("page_break_before") and section_index:
            document.add_page_break()
        add_docx_text(document, section.get("title", ""), style="Heading 1")
        for block in section.get("blocks", []):
            add_docx_block(document, block, variant, manifest_dir, manifest["print_profile"])
    document.save(output)


def main() -> int:
    args = parse_args()
    try:
        manifest_path = args.manifest.expanduser().resolve()
        manifest = load_manifest(manifest_path)
        variants = variants_for(manifest, args.variants)
        output_dir = args.output_dir.expanduser().resolve()
        output_dir.mkdir(parents=True, exist_ok=True)
        prefix = f"{date.today().isoformat()}_{safe_name(manifest['subject'])}_{safe_name(manifest['title'])}"
        labels = {"study": "背诵版", "practice": "练习版", "answers": "解析版"}
        generated: list[str] = []
        for variant in variants:
            stem = f"{prefix}_{labels[variant]}"
            pdf_path = output_dir / f"{stem}.pdf"
            docx_path = output_dir / f"{stem}.docx"
            build_pdf(manifest, variant, pdf_path, manifest_path.parent)
            build_docx(manifest, variant, docx_path, manifest_path.parent)
            generated.extend((str(pdf_path), str(docx_path)))
        print(json.dumps({"generated": generated}, ensure_ascii=False, indent=2))
        return 0
    except (ValueError, RuntimeError, OSError) as exc:
        print(str(exc), file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
